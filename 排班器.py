"""急診 CP-SAT 排班器：直接雙擊此檔案即可啟動本機網頁介面。"""
from __future__ import annotations

import os
import re
import subprocess
import sys
from html import escape
from datetime import date, datetime, timedelta
from io import BytesIO
from pathlib import Path

if os.environ.get("ER_SCHEDULER_LOCAL_AGENT") == "1":
    # GitHub Pages 的本機求解器只使用下方 CP-SAT／匯出函式，不需要載入 Streamlit。
    st = None
    def get_script_run_ctx(): return None
else:
    import streamlit as st
    from streamlit.runtime.scriptrunner import get_script_run_ctx
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill
from ortools.sat.python import cp_model

SHIFTS = ("D", "N")


def month_days(year: int, month: int) -> list[date]:
    first = date(year, month, 1)
    next_month = date(year + (month == 12), month % 12 + 1, 1)
    return [first + timedelta(days=i) for i in range((next_month - first).days)]


def parse_date(value: str) -> date:
    return datetime.strptime(value, "%Y-%m-%d").date()


def has_long_leave(leaves: set[str]) -> bool:
    run, previous = 0, None
    for current in sorted(map(parse_date, leaves)):
        run = run + 1 if previous and current == previous + timedelta(days=1) else 1
        if run >= 7:
            return True
        previous = current
    return False


def solve(payload: dict) -> dict:
    days = month_days(payload["year"], payload["month"])
    keys = [d.isoformat() for d in days]
    index = {key: i for i, key in enumerate(keys)}
    holidays = set(payload["holidays"])
    doctors = payload["doctors"]
    names = [d["name"] for d in doctors]
    by_name = {d["name"]: d for d in doctors}
    errors = []
    if len(names) != len(set(names)):
        errors.append("醫師姓名不可重複")
    if sum(int(d["target_shifts"]) for d in doctors) != len(days) * 2:
        errors.append("所有醫師目標班數加總必須等於當月天數 × 2")
    for d in doctors:
        if d["type"] == "support":
            choices = {(x["date"], x["shift"]) for x in d.get("weekday_availability", [])}
            required = int(d["target_shifts"]) - int(d.get("holiday_quota", 0))
            if len(choices) < required:
                errors.append(f"支援醫師 {d['name']} 需排 {required} 個平日班，但只填 {len(choices)} 個可上平日班")
    fixed = payload["fixed_shifts"]
    occupied = {}
    for item in fixed:
        key = (item["date"], item["shift"])
        if item["doctor"] not in by_name:
            errors.append(f"指定班醫師不存在：{item['doctor']}")
        elif item["date"] not in index or item["shift"] not in SHIFTS:
            errors.append(f"指定班日期或班別無效：{item}")
        elif key in occupied and occupied[key] != item["doctor"]:
            errors.append(f"同一班被兩位醫師指定：{item['date']} {item['shift']}")
        else:
            occupied[key] = item["doctor"]
    if errors:
        return {"status": "INVALID_INPUT", "warnings": errors}

    model = cp_model.CpModel()
    x = {(p, i, s): model.NewBoolVar(f"x_{p}_{i}_{s}") for p in names for i in range(len(days)) for s in SHIFTS}
    work = {(p, i): model.NewBoolVar(f"work_{p}_{i}") for p in names for i in range(len(days))}
    doubles = {(p, i): model.NewBoolVar(f"double_{p}_{i}") for p in names for i in range(len(days))}
    for i, key in enumerate(keys):
        for s in SHIFTS:
            model.AddExactlyOne(x[p, i, s] for p in names)
        for p in names:
            model.AddMaxEquality(work[p, i], [x[p, i, "D"], x[p, i, "N"]])
            model.Add(doubles[p, i] <= x[p, i, "D"])
            model.Add(doubles[p, i] <= x[p, i, "N"])
            model.Add(doubles[p, i] >= x[p, i, "D"] + x[p, i, "N"] - 1)
            doctor = by_name[p]
            if key in set(doctor.get("leave_dates", [])):
                model.Add(x[p, i, "D"] == 0); model.Add(x[p, i, "N"] == 0)
            if doctor["type"] == "support" and key not in holidays:
                allowed = {(a["date"], a["shift"]) for a in doctor.get("weekday_availability", [])}
                for s in SHIFTS:
                    if (key, s) not in allowed:
                        model.Add(x[p, i, s] == 0)
    for item in fixed:
        model.Add(x[item["doctor"], index[item["date"]], item["shift"]] == 1)
    for p in names:
        d = by_name[p]
        model.Add(sum(x[p, i, s] for i in range(len(days)) for s in SHIFTS) == int(d["target_shifts"]))
        for i in range(len(days) - 1):
            model.Add(x[p, i, "N"] + x[p, i + 1, "D"] <= 1)
        if d["type"] == "er":
            for i in range(len(days) - 5):
                model.Add(sum(work[p, j] for j in range(i, i + 6)) <= 5)
        if d["type"] == "support":
            model.Add(sum(x[p, i, s] for i, key in enumerate(keys) if key in holidays for s in SHIFTS) == int(d["holiday_quota"]))
    ers = [d["name"] for d in doctors if d["type"] == "er"]
    if ers:
        support_holidays = sum(int(d.get("holiday_quota", 0)) for d in doctors if d["type"] == "support")
        external_holidays = sum(1 for f in fixed if by_name[f["doctor"]]["type"] == "external" and f["date"] in holidays)
        er_slots = len(holidays) * 2 - support_holidays - external_holidays
        if er_slots < 0:
            return {"status": "INVALID_INPUT", "warnings": ["支援醫師假日配額超過可用假日班數"]}
        low, remainder = divmod(er_slots, len(ers))
        for pos, p in enumerate(ers):
            target = low if pos < len(ers) - remainder else low + 1
            model.Add(sum(x[p, i, s] for i, key in enumerate(keys) if key in holidays for s in SHIFTS) == target)

    objective = []
    double_shift_penalty = int(payload.get("double_shift_penalty", 100))
    for p in names:
        # 有連休 7 天以上者只放寬本人；其他醫師依容忍度套用正／負權重。
        penalty = 0 if has_long_leave(set(by_name[p].get("leave_dates", []))) else double_shift_penalty
        if penalty != 0:
            objective.extend(penalty * doubles[p, i] for i in range(len(days)))
    # 容忍度的軟目標：避免所有正懲罰檔位都退化成「盡量 0 個 24 小時班」。
    desired_doubles = payload.get("desired_double_shifts")
    desired_weight = int(payload.get("desired_double_shifts_weight", 0))
    if desired_doubles is not None and desired_weight > 0:
        total_doubles = model.NewIntVar(0, len(days) * 2, "total_doubles")
        model.Add(total_doubles == sum(doubles.values()))
        difference = model.NewIntVar(0, len(days) * 2, "double_target_difference")
        model.AddAbsEquality(difference, total_doubles - int(desired_doubles))
        objective.append(desired_weight * difference)

    for p in names:
        for i in range(len(days)):
            left_off = 1 if i == 0 else 1 - work[p, i - 1]
            right_off = 1 if i == len(days) - 1 else 1 - work[p, i + 1]
            isolated = model.NewBoolVar(f"one_{p}_{i}")
            model.Add(isolated <= work[p, i]); model.Add(isolated <= left_off); model.Add(isolated <= right_off)
            model.Add(isolated >= work[p, i] + left_off + right_off - 2)
            objective.append(120 * isolated)
            if i < len(days) - 1:
                end_off = 1 if i + 2 == len(days) else 1 - work[p, i + 2]
                two = model.NewBoolVar(f"two_{p}_{i}")
                model.Add(two <= work[p, i]); model.Add(two <= work[p, i + 1]); model.Add(two <= left_off); model.Add(two <= end_off)
                model.Add(two >= work[p, i] + work[p, i + 1] + left_off + end_off - 3)
                objective.append(80 * two)
            if i <= len(days) - 5:
                end_off = 1 if i + 5 == len(days) else 1 - work[p, i + 5]
                five = model.NewBoolVar(f"five_{p}_{i}")
                for j in range(i, i + 5): model.Add(five <= work[p, j])
                model.Add(five <= left_off); model.Add(five <= end_off)
                model.Add(five >= sum(work[p, j] for j in range(i, i + 5)) + left_off + end_off - 6)
                objective.append(50 * five)
    nights = []
    for p in ers:
        count = model.NewIntVar(0, len(days), f"nights_{p}")
        model.Add(count == sum(x[p, i, "N"] for i in range(len(days))))
        nights.append(count)
    if len(nights) > 1:
        lo, hi = model.NewIntVar(0, len(days), "min_nights"), model.NewIntVar(0, len(days), "max_nights")
        model.AddMinEquality(lo, nights); model.AddMaxEquality(hi, nights); objective.append(10 * (hi - lo))
    objective_sum = sum(objective)
    # 「換一個最佳班表」會固定原目標值，並要求至少有一個班與既有方案不同。
    if payload.get("target_objective") is not None:
        model.Add(objective_sum == int(round(payload["target_objective"])))
    for prior_schedule in payload.get("excluded_schedules", []):
        same_assignments = []
        for row in prior_schedule:
            i = index.get(row.get("date"))
            if i is None:
                continue
            if row.get("D") in by_name:
                same_assignments.append(x[row["D"], i, "D"])
            if row.get("N") in by_name:
                same_assignments.append(x[row["N"], i, "N"])
        if same_assignments:
            model.Add(sum(same_assignments) <= len(same_assignments) - 1)
    model.Minimize(objective_sum)
    solver = cp_model.CpSolver(); solver.parameters.max_time_in_seconds = 60; solver.parameters.num_search_workers = 8
    result = solver.Solve(model)
    if result not in (cp_model.OPTIMAL, cp_model.FEASIBLE):
        return {"status": "INFEASIBLE", "warnings": ["在目前所有硬性規則下找不到可行班表。請檢查指定班、預假、目標班數與假日配額。"]}
    schedule = [{"date": key, "D": next(p for p in names if solver.Value(x[p, i, "D"])), "N": next(p for p in names if solver.Value(x[p, i, "N"]))} for i, key in enumerate(keys)]
    return {"status": "OPTIMAL" if result == cp_model.OPTIMAL else "FEASIBLE", "schedule": schedule, "objective": solver.ObjectiveValue(), "solve_time_seconds": solver.WallTime()}


def parse_work(text: str) -> list[dict]:
    result = []
    for day, label in re.findall(r"(\d{4}-\d{2}-\d{2})\s*\(([^)]+)\)", text or ""):
        shift = "D" if "白" in label or label.strip().upper() in {"D", "DAY"} else "N" if "夜" in label or label.strip().upper() in {"N", "NIGHT"} else None
        if shift: result.append({"date": day, "shift": shift})
    return result


def cell_date(value) -> str | None:
    if isinstance(value, datetime): return value.date().isoformat()
    if isinstance(value, date): return value.isoformat()
    for fmt in ("%Y-%m-%d", "%Y/%m/%d"):
        try: return datetime.strptime(str(value or "")[:10], fmt).date().isoformat()
        except ValueError: pass
    return None


def make_payload(uploaded) -> dict:
    wb = load_workbook(uploaded, data_only=True)
    required = ["表單回應", "急專醫師名單", "支援醫師名單", "特殊日期設定"]
    missing = [name for name in required if name not in wb.sheetnames]
    if missing: raise ValueError("缺少工作表：" + "、".join(missing))
    rows = lambda name: [["" if v is None else str(v).strip() for v in r] for r in wb[name].iter_rows(values_only=True)]
    er_rows, support_rows, response_rows = rows("急專醫師名單")[1:], rows("支援醫師名單")[1:], rows("表單回應")[1:]
    months = set()
    for r in response_rows:
        for value in r[2:4]: months.update(re.findall(r"(\d{4})-(\d{2})-\d{2}", value))
    if len(months) != 1: raise ValueError("無法從預班／預假判斷唯一的排班月份")
    year, month = map(int, next(iter(months))); prefix = f"{year}-{month:02d}-"
    er = {r[0]: int(float(r[1])) if len(r) > 1 and r[1] else 16 for r in er_rows if r and r[0]}
    support = [r[0] for r in support_rows if r and r[0]]
    if not support: raise ValueError("支援醫師名單不可空白")
    marker = next((i for i, r in enumerate(support_rows) if r and r[0] and len(r) > 1 and r[1].upper() == "V"), -1)
    support_set = set(support); prefs = {}
    for r in response_rows:
        if len(r) < 4 or not r[1]: continue
        work = [x for x in parse_work(r[2]) if x["date"].startswith(prefix)]
        leaves = set(re.findall(r"\d{4}-\d{2}-\d{2}", r[3] or ""))
        prefs[r[1]] = {"work": work, "leaves": {x for x in leaves if x.startswith(prefix)}}
    holidays, makeup = set(), set()
    for r in list(wb["特殊日期設定"].iter_rows(values_only=True))[1:]:
        if len(r) < 2: continue
        day, label = cell_date(r[0]), str(r[1] or "")
        if not day or not day.startswith(prefix): continue
        if "補" in label or "上班" in label: makeup.add(day)
        elif "假" in label or "休" in label: holidays.add(day)
    for d in month_days(year, month):
        if d.weekday() >= 5: holidays.add(d.isoformat())
    holidays -= makeup
    doctors, fixed = [], []
    for name, target in er.items():
        pref = prefs.get(name, {"work": [], "leaves": set()})
        doctors.append({"name": name, "type": "er", "target_shifts": target, "leave_dates": sorted(pref["leaves"])})
        fixed += [{"doctor": name, **x} for x in pref["work"]]
    for name, pref in prefs.items():
        if name in er or name in support_set: continue
        doctors.append({"name": name, "type": "external", "target_shifts": len(pref["work"]), "leave_dates": sorted(pref["leaves"])})
        fixed += [{"doctor": name, **x} for x in pref["work"]]
    total = len(month_days(year, month)) * 2 - sum(d["target_shifts"] for d in doctors)
    if total < 0: raise ValueError("急專與外援的目標班數已超過本月總班數")
    external_holidays = sum(1 for f in fixed if f["doctor"] not in er and f["doctor"] not in support_set and f["date"] in holidays)
    er_holiday_target = (len(holidays) * 2 * 3 + 3) // 4
    support_holidays = max(0, len(holidays) * 2 - er_holiday_target - external_holidays)
    base, remainder = divmod(total, len(support)); holiday_quota = {name: 0 for name in support}
    for offset in range(support_holidays): holiday_quota[support[(marker + 1 + offset) % len(support)]] += 1
    for pos, name in enumerate(support):
        pref = prefs.get(name, {"work": [], "leaves": set()})
        doctors.append({"name": name, "type": "support", "target_shifts": base + (pos < remainder), "holiday_quota": holiday_quota[name], "leave_dates": sorted(pref["leaves"]), "weekday_availability": [x for x in pref["work"] if x["date"] not in holidays]})
    return {"year": year, "month": month, "holidays": sorted(holidays), "doctors": doctors, "fixed_shifts": fixed}


def preview_rows(payload: dict) -> list[dict]:
    return [{"姓名": d["name"], "類別": {"er": "急專", "external": "外援", "support": "支援"}[d["type"]], "目標班數": d["target_shifts"], "平日": d["target_shifts"] - d.get("holiday_quota", 0), "假日": d.get("holiday_quota", 0)} for d in payload["doctors"]]


def validate_result(payload: dict, result: dict) -> tuple[list[dict], list[dict]]:
    """回傳每人統計與所有硬規則違規；正常最佳解應為零違規。"""
    schedule = result["schedule"]
    holidays, by_name = set(payload["holidays"]), {d["name"]: d for d in payload["doctors"]}
    assignment = {name: {} for name in by_name}
    for row in schedule:
        assignment[row["D"]].setdefault(row["date"], []).append("白")
        assignment[row["N"]].setdefault(row["date"], []).append("夜")
    violations, stats = [], []
    fixed = {(x["doctor"], x["date"], x["shift"]) for x in payload["fixed_shifts"]}
    for name, doctor in by_name.items():
        items = assignment[name]
        total = sum(len(v) for v in items.values())
        day_count = sum(v.count("白") for v in items.values())
        night_count = sum(v.count("夜") for v in items.values())
        holiday_count = sum(len(v) for key, v in items.items() if key in holidays)
        doubles = sum(len(v) == 2 for v in items.values())
        stats.append({"姓名": name, "類別": {"er": "急專", "external": "外援", "support": "支援"}[doctor["type"]], "目標": doctor["target_shifts"], "實排": total, "白班": day_count, "夜班": night_count, "假日": holiday_count, "24 小時班": doubles, "結果": "通過" if total == doctor["target_shifts"] else "班數不符"})
        if total != doctor["target_shifts"]: violations.append({"類型": "目標班數", "醫師": name, "日期": "", "說明": f"目標 {doctor['target_shifts']}，實排 {total}"})
        if doctor["type"] == "support" and holiday_count != doctor["holiday_quota"]:
            violations.append({"類型": "支援假日配額", "醫師": name, "日期": "", "說明": f"配額 {doctor['holiday_quota']}，實排 {holiday_count}"})
        if doctor["type"] == "support":
            allowed = {(x["date"], "白" if x["shift"] == "D" else "夜") for x in doctor.get("weekday_availability", [])}
            for day, shifts in items.items():
                for shift in shifts:
                    if day not in holidays and (day, shift) not in allowed:
                        violations.append({"類型": "支援平日可上班時段", "醫師": name, "日期": day, "說明": f"未填可上{shift}班"})
        for leave in doctor.get("leave_dates", []):
            if leave in items: violations.append({"類型": "預假", "醫師": name, "日期": leave, "說明": "預假日被排班"})
        dates = sorted(items)
        for current in dates:
            next_day = (parse_date(current) + timedelta(days=1)).isoformat()
            if "夜" in items[current] and "白" in items.get(next_day, []):
                violations.append({"類型": "夜接白", "醫師": name, "日期": current, "說明": "夜班後隔日排白班"})
        if doctor["type"] == "er":
            work_dates = set(dates)
            for start in month_days(payload["year"], payload["month"])[:-5]:
                window = [(start + timedelta(days=x)).isoformat() for x in range(6)]
                if sum(d in work_dates for d in window) > 5:
                    violations.append({"類型": "急專六日規則", "醫師": name, "日期": start.isoformat(), "說明": "連續六日工作超過五日"})
    for doctor, day, shift in fixed:
        label = "白" if shift == "D" else "夜"
        if label not in assignment[doctor].get(day, []):
            violations.append({"類型": "指定班", "醫師": doctor, "日期": day, "說明": f"指定{label}班未排入"})
    ers = [d["name"] for d in payload["doctors"] if d["type"] == "er"]
    if ers:
        support_holidays = sum(d.get("holiday_quota", 0) for d in payload["doctors"] if d["type"] == "support")
        external_holidays = sum(1 for d, day, _ in fixed if by_name[d]["type"] == "external" and day in holidays)
        low, remainder = divmod(len(holidays) * 2 - support_holidays - external_holidays, len(ers))
        for pos, name in enumerate(ers):
            target = low if pos < len(ers) - remainder else low + 1
            actual = next(row["假日"] for row in stats if row["姓名"] == name)
            if actual != target:
                violations.append({"類型": "急專假日均衡", "醫師": name, "日期": "", "說明": f"目標 {target}，實排 {actual}"})
    return stats, violations


def calendar_html(payload: dict, result: dict, selected_doctor: str | None = None) -> str:
    """建立與人工審核習慣一致的週曆：日期、白班、夜班各一列。"""
    year, month = payload["year"], payload["month"]
    by_date = {row["date"]: row for row in result["schedule"]}
    first = date(year, month, 1)
    start = first - timedelta(days=first.weekday())
    last = month_days(year, month)[-1]
    end = last + timedelta(days=6 - last.weekday())
    rows = ["<table class='schedule-calendar'><thead><tr><th></th>"]
    for label in ("一", "二", "三", "四", "五", "六", "日"):
        rows.append(f"<th class='{'weekend' if label in {'六', '日'} else ''}'>{label}</th>")
    rows.append("</tr></thead><tbody>")
    cursor = start
    while cursor <= end:
        week = [cursor + timedelta(days=i) for i in range(7)]
        rows.append("<tr class='date-row'><th></th>")
        for day in week:
            text = str(day.day) if day.month == month else ""
            rows.append(f"<td>{text}</td>")
        rows.append("</tr><tr><th>白班</th>")
        for day in week:
            name = by_date.get(day.isoformat(), {}).get("D", "") if day.month == month else ""
            style = " class='doctor-highlight'" if name == selected_doctor else ""
            rows.append(f"<td{style}>{escape(name)}</td>")
        rows.append("</tr><tr><th>夜班</th>")
        for day in week:
            name = by_date.get(day.isoformat(), {}).get("N", "") if day.month == month else ""
            style = " class='doctor-highlight'" if name == selected_doctor else ""
            rows.append(f"<td{style}>{escape(name)}</td>")
        rows.append("</tr>")
        cursor += timedelta(days=7)
    rows.append("</tbody></table>")
    return "".join(rows)


def render_clickable_calendar(payload: dict, result: dict) -> None:
    """月曆中的醫師姓名本身可點選，點後標亮該醫師的所有班別。"""
    year, month = payload["year"], payload["month"]
    by_date = {row["date"]: row for row in result["schedule"]}
    selected = st.session_state.get("highlight_doctor")
    first, last = date(year, month, 1), month_days(year, month)[-1]
    start, end = first - timedelta(days=first.weekday()), last + timedelta(days=6 - last.weekday())
    st.markdown("""
    <style>
    .calendar-label, .calendar-day, .calendar-date {text-align:center; min-height:34px; padding:5px 2px; border:1px solid #9aa8b6; margin-bottom:-1px;}
    .calendar-label {font-weight:700; background:#f7f9fb;}
    .calendar-day {font-weight:700; color:white; background:#4f81bd;}
    .calendar-date {font-weight:700; background:#dce6f1;}
    .calendar-empty {background:#e7e6e6;}
    div[data-testid="stHorizontalBlock"] button[kind="secondary"] {min-height:34px; padding:2px 3px;}
    </style>
    """, unsafe_allow_html=True)
    cursor, week_no = start, 0
    labels = ["一", "二", "三", "四", "五", "六", "日"]
    while cursor <= end:
        week = [cursor + timedelta(days=i) for i in range(7)]
        widths = [0.65] + [1] * 7
        header = st.columns(widths)
        header[0].markdown("<div class='calendar-day'></div>", unsafe_allow_html=True)
        for col, day in enumerate(week):
            content = labels[col] if day.month == month else ""
            header[col + 1].markdown(f"<div class='calendar-day'>{content}</div>", unsafe_allow_html=True)
        dates = st.columns(widths)
        dates[0].markdown("<div class='calendar-date'></div>", unsafe_allow_html=True)
        for col, day in enumerate(week):
            text = str(day.day) if day.month == month else ""
            extra = " calendar-empty" if day.month != month else ""
            dates[col + 1].markdown(f"<div class='calendar-date{extra}'>{text}</div>", unsafe_allow_html=True)
        for shift, label in (("D", "白班"), ("N", "夜班")):
            row = st.columns(widths)
            row[0].markdown(f"<div class='calendar-label'>{label}</div>", unsafe_allow_html=True)
            for col, day in enumerate(week):
                if day.month != month:
                    row[col + 1].markdown("<div class='calendar-label calendar-empty'></div>", unsafe_allow_html=True)
                    continue
                name = by_date[day.isoformat()][shift]
                if row[col + 1].button(name, key=f"calendar-{week_no}-{shift}-{day.isoformat()}", use_container_width=True, type="primary" if name == selected else "secondary"):
                    st.session_state["highlight_doctor"] = None if name == selected else name
                    st.rerun()
        cursor += timedelta(days=7); week_no += 1


def select_highlight_doctor(name: str) -> None:
    """按鈕 callback：必須在 radio widget 建立前更新其選取狀態。"""
    st.session_state["highlight_doctor"] = None if name == st.session_state.get("highlight_doctor") else name
    st.session_state["result_view"] = "醫師月曆"


def change_active_schedule(delta: int, view: str) -> None:
    options = st.session_state["schedule_options"]
    current = st.session_state.get("active_schedule_index", 0)
    target = max(0, min(len(options) - 1, current + delta))
    st.session_state["active_schedule_index"] = target
    st.session_state["schedule_result"] = options[target]
    st.session_state["result_view"] = view


def return_to_tolerance_settings() -> None:
    """保留原始 Excel，回到班數預覽與 24 小時班容忍度設定。"""
    st.session_state.pop("schedule_result", None)
    st.session_state.pop("highlight_doctor", None)


def find_alternative_schedule(view: str) -> None:
    payload = st.session_state["schedule_payload"]
    alternative_payload = dict(payload)
    alternative_payload["target_objective"] = st.session_state["schedule_best_objective"]
    alternative_payload["excluded_schedules"] = st.session_state["schedule_history"]
    alternative = solve(alternative_payload)
    st.session_state["result_view"] = view
    if alternative["status"] not in {"OPTIMAL", "FEASIBLE"}:
        st.session_state["alternative_warning"] = "已找不到另一個具有相同目標值的班表。"
        return
    st.session_state["schedule_result"] = alternative
    st.session_state["schedule_history"].append(alternative["schedule"])
    st.session_state["schedule_options"].append(alternative)
    st.session_state["active_schedule_index"] = len(st.session_state["schedule_options"]) - 1


def collect_best_schedules(payload: dict, maximum: int = 20, progress=None) -> tuple[list[dict], str | None]:
    """先求最佳目標值，再列舉最多 maximum 組具有相同目標值的不同班表。"""
    if progress:
        progress(0, maximum, "正在尋找第一組最佳班表…")
    first = solve(payload)
    if first["status"] not in {"OPTIMAL", "FEASIBLE"}:
        return [first], None
    # 時限內僅得到可行解時，不把它誤稱為「所有最佳解」。
    if first["status"] != "OPTIMAL":
        return [first], "首次搜尋在時限內僅得到 FEASIBLE；已保留目前最佳班表，未列舉替代方案。"

    options = [first]
    history = [first["schedule"]]
    while len(options) < maximum:
        if progress:
            progress(len(options), maximum, f"已找到 {len(options)} 組最佳班表，正在找下一組…")
        alternative_payload = dict(payload)
        alternative_payload["target_objective"] = first["objective"]
        alternative_payload["excluded_schedules"] = history
        alternative = solve(alternative_payload)
        if alternative["status"] not in {"OPTIMAL", "FEASIBLE"}:
            break
        options.append(alternative)
        history.append(alternative["schedule"])

    message = None
    if len(options) < maximum:
        message = f"已找完所有可找到的同品質最佳班表，共 {len(options)} 組。"
    return options, message


def render_schedule_navigation(active: int, options: list[dict], view: str) -> None:
    left, label, right = st.columns([1, 4, 1])
    left.button("<", key=f"solution-prev-{view}", disabled=active == 0, use_container_width=True, on_click=change_active_schedule, args=(-1, view))
    label.markdown(f"<div style='text-align:center; padding:0.45rem; font-weight:700;'>現在方案 {active + 1} / {len(options)}｜目標值 {options[active]['objective']:.0f}</div>", unsafe_allow_html=True)
    right.button(">", key=f"solution-next-{view}", disabled=active == len(options) - 1, use_container_width=True, on_click=change_active_schedule, args=(1, view))


def export_workbook(original: bytes, result: dict, payload: dict) -> bytes:
    wb = load_workbook(BytesIO(original))
    if "CP-SAT 班表" in wb.sheetnames: del wb["CP-SAT 班表"]
    ws = wb.create_sheet("CP-SAT 班表")
    ws.append(["CP-SAT 排班結果", result["status"], f"目標值：{result.get('objective', '')}"])
    ws.append(["日期", "白班", "夜班"])
    for row in result.get("schedule", []): ws.append([row["date"], row["D"], row["N"]])
    ws["A1"].font = Font(bold=True, color="FFFFFF"); ws["A1"].fill = PatternFill("solid", fgColor="1F4E78")
    for cell in ws[2]: cell.font = Font(bold=True, color="FFFFFF"); cell.fill = PatternFill("solid", fgColor="4F81BD")
    ws.column_dimensions["A"].width = 16; ws.column_dimensions["B"].width = 18; ws.column_dimensions["C"].width = 18; ws.freeze_panes = "A3"
    output = BytesIO(); wb.save(output); return output.getvalue()


def _word_cell_fill(cell, color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    shading = OxmlElement("w:shd"); shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _word_font(run, name: str, size: int, bold: bool = False, color: str | None = None) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    run.font.name = name; run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor.from_string(color)


def export_docx(payload: dict, result: dict) -> bytes:
    """依原 Google 文件版型建立可下載的 Word 值班表。"""
    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ROW_HEIGHT_RULE
    from docx.shared import Cm, Mm, Pt

    year, month = payload["year"], payload["month"]
    roc_year, days = year - 1911, month_days(year, month)
    by_date, holidays = {r["date"]: r for r in result["schedule"]}, set(payload["holidays"])
    doc = Document(); section = doc.sections[0]
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Cm(1); section.bottom_margin = Cm(1); section.left_margin = Cm(1); section.right_margin = Cm(1)
    from docx.oxml.ns import qn
    normal = doc.styles["Normal"]; normal.font.name = "DFKai-SB"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "DFKai-SB")
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after = Pt(0); title.paragraph_format.line_spacing = 1.8
    _word_font(title.add_run("國軍花蓮總醫院急診醫學科"), "Times New Roman", 18, True)
    subtitle = doc.add_paragraph(); subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER; subtitle.paragraph_format.space_after = Pt(6); subtitle.paragraph_format.line_spacing = 1.8
    _word_font(subtitle.add_run(f"中華民國{roc_year}年 {month} 月值班表"), "Times New Roman", 16, True)
    first = days[0]; start = first - timedelta(days=first.weekday()); end = days[-1] + timedelta(days=6 - days[-1].weekday())
    weeks = ((end - start).days + 1) // 7
    table = doc.add_table(rows=weeks * 4, cols=8); table.style = "Table Grid"; table.autofit = False
    label_width, day_width = Cm(1.15), Cm(2.53)
    weekday_names = ["一", "二", "三", "四", "五", "六", "日"]
    for week in range(weeks):
        base, monday = week * 4, start + timedelta(days=week * 7)
        # 首欄以兩列為一組合併：日期區留白，白／夜班區直排顯示「值班醫師」。
        table.cell(base, 0).merge(table.cell(base + 1, 0))
        duty_label = table.cell(base + 2, 0).merge(table.cell(base + 3, 0))
        duty_label.text = "值\n班\n醫\n師"
        # 白、夜班列各 1.1 cm；合併後的「值班醫師」區域為 2.2 cm。
        duty_row_height = Cm(1.1)
        for row_offset, height in enumerate((Cm(0.48), Cm(0.5), duty_row_height, duty_row_height)):
            row = table.rows[base + row_offset]; row.height = height
            row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY if row_offset in (2, 3) else WD_ROW_HEIGHT_RULE.AT_LEAST
        for col in range(8):
            for row_offset in range(4):
                cell = table.cell(base + row_offset, col); cell.width = label_width if col == 0 else day_width
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        duty_label.paragraphs[0].paragraph_format.space_before = Pt(0)
        duty_label.paragraphs[0].paragraph_format.space_after = Pt(0)
        duty_label.paragraphs[0].paragraph_format.line_spacing = 1.0
        for run in duty_label.paragraphs[0].runs: _word_font(run, "Times New Roman", 12, True)
        for col in range(7):
            current = monday + timedelta(days=col); cell_col = col + 1
            if current.month != month:
                for row_offset in range(4): _word_cell_fill(table.cell(base + row_offset, cell_col), "E7E6E6")
                continue
            row = by_date[current.isoformat()]
            _word_font(table.cell(base, cell_col).paragraphs[0].add_run(weekday_names[col]), "DFKai-SB", 14)
            _word_font(table.cell(base + 1, cell_col).paragraphs[0].add_run(str(current.day)), "DFKai-SB", 14)
            _word_font(table.cell(base + 2, cell_col).paragraphs[0].add_run(row["D"]), "DFKai-SB", 14)
            _word_font(table.cell(base + 3, cell_col).paragraphs[0].add_run(row["N"]), "DFKai-SB", 14)
            if current.isoformat() in holidays:
                _word_cell_fill(table.cell(base, cell_col), "FFFF00"); _word_cell_fill(table.cell(base + 1, cell_col), "FFFF00")
    weekday_shifts = sum(2 for day in days if day.isoformat() not in holidays)
    holiday_shifts = len(days) * 2 - weekday_shifts
    blank = doc.add_paragraph(); blank.paragraph_format.space_after = Pt(0); blank.paragraph_format.line_spacing = 1
    note = doc.add_paragraph(); note.paragraph_format.space_after = Pt(2)
    _word_font(note.add_run("注意：每週三醫療部會議完為科務會議。(值勤時段為白班08:00-20:00；夜班20:00-08:00)"), "Times New Roman", 13, False, "FF0000")
    note2 = doc.add_paragraph(); note2.paragraph_format.space_after = Pt(0)
    _word_font(note2.add_run("(1) 班表按照"), "Times New Roman", 13)
    _word_font(note2.add_run("急診室排班規定"), "Times New Roman", 13, True)
    _word_font(note2.add_run("擬定"), "Times New Roman", 13)
    note3 = doc.add_paragraph(); note3.paragraph_format.space_after = Pt(0)
    _word_font(note3.add_run(f"(2) 排定後門急診部主任可依各醫師門診時間實施微調 共 {len(days) * 2} 班（{weekday_shifts}+{holiday_shifts}）"), "Times New Roman", 13)
    output = BytesIO(); doc.save(output); return output.getvalue()


def render() -> None:
    st.set_page_config(page_title="急診排班器", page_icon="📅", layout="wide")
    st.title("急診排班器")
    has_result = "schedule_result" in st.session_state
    if not has_result:
        st.caption("離線 CP-SAT 排班｜上傳 Excel、確認配額、產生班表")
        file = st.file_uploader("選擇 Excel 班表資料", type="xlsx")
        if file:
            raw, source_name = file.getvalue(), file.name
        elif "schedule_raw" in st.session_state:
            # 從結果頁返回時保留同一份資料，不必再次選檔。
            raw, source_name = st.session_state["schedule_raw"], st.session_state.get("schedule_name", "班表資料.xlsx")
        else:
            return
        try:
            payload = make_payload(BytesIO(raw)); rows = preview_rows(payload)
        except Exception as exc:
            st.error(str(exc)); return
        st.markdown("""
        <style>
        .preview-heading {font-size: 2.0rem; font-weight: 700; margin: 1.25rem 0 0.65rem;}
        div[data-testid="stMetricLabel"] {font-size: 1.05rem;}
        div[data-testid="stMetricValue"] {font-size: 2.15rem;}
        </style>
        """, unsafe_allow_html=True)
        st.markdown(f"<div class='preview-heading'>{payload['year']} 年 {payload['month']} 月班數分配預覽</div>", unsafe_allow_html=True)
        a, b, c, d = st.columns(4)
        a.metric("本月總班數", len(month_days(payload["year"], payload["month"])) * 2)
        b.metric("急專目標班數", sum(r["目標班數"] for r in rows if r["類別"] == "急專"))
        c.metric("外援鎖定班數", sum(r["目標班數"] for r in rows if r["類別"] == "外援"))
        d.metric("支援分配班數", sum(r["目標班數"] for r in rows if r["類別"] == "支援"))
        st.dataframe(rows, use_container_width=True, hide_index=True, height=480)
        tolerance_labels = {
            0: "0：強烈避免",
            25: "25：可接受少量",
            50: "50：容許約 1–2 個",
            75: "75：容許較常出現",
            100: "100：主動偏好用 24 小時班換取較完整休息／長假",
        }
        # 每一段皆有不同的「全月 24 小時班軟目標」，避免 0–50 都只排成 0 班。
        # 長假醫師仍只會放寬他本人的逐班懲罰，並不會影響其他人。
        double_profiles = {
            0: {"penalty": 200, "target": 0, "target_weight": 0},
            25: {"penalty": 15, "target": 1, "target_weight": 50},
            50: {"penalty": 10, "target": 2, "target_weight": 50},
            75: {"penalty": 5, "target": 4, "target_weight": 45},
            100: {"penalty": -5, "target": 6, "target_weight": 35},
        }
        tolerance = st.select_slider(
            "24 小時班容忍度",
            options=[0, 25, 50, 75, 100],
            value=st.session_state.get("schedule_tolerance", 50),
            format_func=lambda value: tolerance_labels[value],
            help="長假 7 天以上者只放寬該醫師；100 會主動偏好 24 小時班。",
        )
        profile = double_profiles[tolerance]
        payload["double_shift_penalty"] = profile["penalty"]
        payload["desired_double_shifts"] = profile["target"]
        payload["desired_double_shifts_weight"] = profile["target_weight"]
        if tolerance == 0:
            st.caption("目前設定：強烈避免 24 小時班。")
        else:
            st.caption(f"目前設定：全月約 {profile['target']} 個 24 小時班的軟目標；硬性規則優先。")
        if st.button("產生 CP-SAT 班表", type="primary", use_container_width=True):
            progress_bar = st.progress(0, text="正在準備最佳班表搜尋…")
            with st.status("正在尋找最佳班表與替代方案…", expanded=True) as status_box:
                def update_progress(found: int, maximum: int, message: str) -> None:
                    progress_bar.progress(int(found / maximum * 100), text=message)
                    status_box.write(message)

                options, search_note = collect_best_schedules(payload, maximum=20, progress=update_progress)
                result = options[0]
                if result["status"] == "OPTIMAL":
                    final_message = f"完成：共找到 {len(options)} 組同品質最佳班表。"
                elif result["status"] == "FEASIBLE":
                    final_message = "完成：已找到目前可行班表。"
                else:
                    final_message = "完成：找不到符合所有硬性規則的班表。"
                progress_bar.progress(100, text=final_message)
                status_box.update(label=final_message, state="complete", expanded=False)
            if result["status"] not in {"OPTIMAL", "FEASIBLE"}:
                st.error("；".join(result.get("warnings", [result["status"]]))); return
            st.session_state["schedule_result"] = result
            st.session_state["schedule_payload"] = payload
            st.session_state["schedule_raw"] = raw
            st.session_state["schedule_name"] = source_name
            st.session_state["schedule_tolerance"] = tolerance
            st.session_state["schedule_history"] = [item["schedule"] for item in options]
            st.session_state["schedule_best_objective"] = result["objective"]
            st.session_state["schedule_options"] = options
            st.session_state["initial_search_note"] = search_note
            st.session_state["active_schedule_index"] = 0
            st.session_state["result_view"] = "醫師月曆"
            st.rerun()
        else:
            return
    result = st.session_state["schedule_result"]
    payload = st.session_state["schedule_payload"]
    raw = st.session_state["schedule_raw"]
    file_name = st.session_state["schedule_name"]
    options = st.session_state.get("schedule_options", [result])
    st.session_state.setdefault("schedule_options", options)
    st.session_state.setdefault("schedule_history", [item["schedule"] for item in options])
    st.session_state.setdefault("schedule_best_objective", options[0]["objective"])
    active = st.session_state.get("active_schedule_index", 0)
    stats, violations = validate_result(payload, result)
    action_col, status_col = st.columns([1, 3])
    with action_col:
        st.button("← 返回調整容忍度", use_container_width=True, on_click=return_to_tolerance_settings)
    with status_col:
        st.success(f"{result['status']}｜目標值 {result['objective']:.0f}｜耗時 {result['solve_time_seconds']:.1f} 秒")
    initial_search_note = st.session_state.pop("initial_search_note", None)
    if initial_search_note:
        st.info(initial_search_note)
    if st.session_state.pop("alternative_warning", None):
        st.warning("已找不到另一個具有相同目標值的班表。")
    view = st.radio("結果檢視", ["醫師月曆", "每日班表", "結果驗證"], horizontal=True, label_visibility="collapsed", key="result_view")
    if view == "每日班表":
        render_schedule_navigation(active, options, view)
        st.dataframe(result["schedule"], use_container_width=True, hide_index=True)
    elif view == "醫師月曆":
        render_schedule_navigation(active, options, view)
        calendar_col, doctor_col = st.columns([5, 1])
        with calendar_col:
            st.markdown("""
            <style>
            .schedule-calendar { width: 100%; border-collapse: collapse; font-size: 16px; table-layout: fixed; }
            .schedule-calendar th, .schedule-calendar td { border: 1px solid #9aa8b6; text-align: center; padding: 4px; height: 31px; }
            .schedule-calendar thead th { background: #4f81bd; color: white; font-weight: 700; height: 34px; }
            .schedule-calendar thead .weekend { background: #3f70ad; }
            .schedule-calendar tbody th { background: #f7f9fb; width: 62px; font-weight: 700; }
            .schedule-calendar .date-row td { background: #dce6f1; font-weight: 700; }
            .schedule-calendar .date-row th { background: #dce6f1; }
            .schedule-calendar td.doctor-highlight { background: #fff2a8; color: #7a2e00; font-weight: 700; }
            </style>
            """, unsafe_allow_html=True)
            st.markdown(calendar_html(payload, result, st.session_state.get("highlight_doctor")), unsafe_allow_html=True)
        with doctor_col:
            st.caption("點選醫師標亮")
            name_cols = st.columns(2)
            for pos, doctor in enumerate(payload["doctors"]):
                name = doctor["name"]
                name_cols[pos % 2].button(
                    name,
                    key=f"highlight-{name}",
                    use_container_width=True,
                    type="primary" if name == st.session_state.get("highlight_doctor") else "secondary",
                    on_click=select_highlight_doctor,
                    args=(name,),
                )
    else:
        render_schedule_navigation(active, options, view)
        a, b, c = st.columns(3)
        a.metric("硬性規則違規", len(violations))
        b.metric("24 小時班", sum(row["24 小時班"] for row in stats))
        c.metric("急專夜班差距", max((row["夜班"] for row in stats if row["類別"] == "急專"), default=0) - min((row["夜班"] for row in stats if row["類別"] == "急專"), default=0))
        if violations:
            st.error("發現硬性規則違規，請勿使用此班表。")
            st.dataframe(violations, use_container_width=True, hide_index=True)
        else:
            st.success("所有已實作硬性規則均通過。")
        st.subheader("急專醫師班數統計")
        st.dataframe([row for row in stats if row["類別"] == "急專"], use_container_width=True, hide_index=True)
    name = Path(file_name).stem + "_CP-SAT.xlsx"
    excel_bytes = st.session_state.get("schedule_excel_bytes")
    if excel_bytes is None or st.session_state.get("export_schedule") is not result:
        excel_bytes = export_workbook(raw, result, payload)
        st.session_state["schedule_excel_bytes"] = excel_bytes
        st.session_state["export_schedule"] = result
    st.download_button("下載班表 Excel", excel_bytes, file_name=name, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", type="primary", use_container_width=True)
    word_name = f"急診大值值班表{payload['year'] - 1911}-{payload['month']:02d}.docx"
    docx_bytes = st.session_state.get("schedule_docx_bytes")
    if docx_bytes is None or st.session_state.get("export_docx_schedule") is not result:
        docx_bytes = export_docx(payload, result)
        st.session_state["schedule_docx_bytes"] = docx_bytes
        st.session_state["export_docx_schedule"] = result
    st.download_button("下載 Word 值班表", docx_bytes, file_name=word_name, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)


if __name__ == "__main__":
    if get_script_run_ctx() is not None:
        render()
    else:
        subprocess.run(
            [sys.executable, "-m", "streamlit", "run", str(Path(__file__).resolve()), "--server.headless=false"],
            env=os.environ.copy(),
            check=False,
        )
